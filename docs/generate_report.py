"""Генератор docx-отчёта по выполненным заданиям 4.x, 5.x, 6.x.

Запуск: python docs/generate_report.py
Результат: docs/report.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _add_section(doc: Document, title: str, points: list[tuple[str, str, list[str]]]) -> None:
    _add_heading(doc, title, level=1)
    for code, summary, files in points:
        _add_heading(doc, code, level=2)
        p = doc.add_paragraph()
        p.add_run("Что сделано: ").bold = True
        p.add_run(summary)
        if files:
            p = doc.add_paragraph()
            p.add_run("Где смотреть: ").bold = True
            p.add_run("; ".join(files))


def main() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Elevator Control System — отчёт по заданиям 4–6", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        "Отчёт описывает попунктное выполнение требований к проекту: CQRS, BFF, "
        "Domain Events с очередью, eventual consistency, Git-workflow, observability, "
        "кэш и блок безопасности (валидация, JWT, защита от атак)."
    ).italic = True

    # ---------- Задание 4 ----------
    _add_section(
        doc,
        "Задание 4. CQRS / BFF / Events / Eventual Consistency",
        [
            (
                "4.1.1 CQRS — разделение command/query (один эндпоинт для read/write запрещён)",
                "Введены отдельные сервисы для команд и запросов. Команды (POST/PATCH/DELETE) "
                "идут в *CommandService и работают через ORM, события доменной модели публикуются. "
                "Запросы (GET) идут в *QueryService и читают денормализованную read-модель через raw SQL.",
                [
                    "elevator_control/application/services.py — *CommandService",
                    "elevator_control/application/queries/*.py — *QueryService",
                    "elevator_control/adapters/inbound/api/v1/lifts.py и др. — GET в Query, write в Command",
                    "elevator_control/adapters/inbound/api/deps.py — раздельные DI: *CmdDep / *QueryDep",
                ],
            ),
            (
                "4.1.2 Разные модели данных: write — нормализованная, read — денормализованная",
                "Write-модель — исходные таблицы lifts/events/service_requests/sensors/technicians/reports "
                "(нормализованы, миграции 001/002). Read-модель — таблицы *_read с агрегатами и подсказками "
                "для UI (sensors_count, open_events_count, max_sensor_ratio, technician_name и т.д.). "
                "Создаётся миграцией 003.",
                [
                    "alembic/versions/003_cqrs_read_model.py — DDL read-таблиц + первичный backfill",
                    "elevator_control/application/queries/*_queries.py — DTO и SELECT по *_read",
                    "elevator_control/application/read_sync.py — UPSERT, пересоберающий строку read-модели",
                ],
            ),
            (
                "4.1.3 Query без ORM",
                "Read-сторона использует прямой пул asyncpg. Все запросы — параметризованные ($1, $2, ...), "
                "без SQLAlchemy ORM. Это закрывает риск SQL-инъекций (см. 6.3.2) и позволяет писать тонкие "
                "denormalised SELECT под конкретный экран UI.",
                [
                    "elevator_control/infrastructure/raw_pool.py — asyncpg.create_pool",
                    "elevator_control/application/queries/base.py — _fetch_all/_fetch_one/_fetch_value",
                    "elevator_control/application/queries/*.py — конкретные запросы",
                ],
            ),
            (
                "4.2.1 У каждого клиента свой backend (BFF)",
                "Под каждый клиент выделен отдельный набор эндпоинтов: /bff/web (SPA), /bff/mobile (моб. клиент), "
                "/bff/desktop (pywebview). Каждый BFF использует только Query-сторону и НЕ дублирует write-логику.",
                [
                    "elevator_control/adapters/inbound/api/bff/__init__.py — корневой роутер /bff",
                    "elevator_control/adapters/inbound/api/bff/web.py / mobile.py / desktop.py",
                ],
            ),
            (
                "4.2.2 Разные DTO для разных клиентов",
                "Web получает компактный дашборд с risk_score; Mobile — карточки с предсобранным title и "
                "светофорным status (ok/warning/critical); Desktop — «рабочее место по лифту» (лифт + датчики "
                "+ события + заявки в одном ответе).",
                [
                    "elevator_control/adapters/inbound/api/bff/schemas.py — Web/Mobile/Desktop DTO",
                ],
            ),
            (
                "4.2.3 Агрегация данных в BFF",
                "Каждый BFF-эндпоинт параллельно (asyncio.gather) дёргает несколько Query-сервисов и собирает "
                "результат под нужды конкретного UI. Один HTTP-вызов с фронта = один ответ с уже подготовленными данными.",
                [
                    "elevator_control/adapters/inbound/api/bff/web.py::web_dashboard",
                    "elevator_control/adapters/inbound/api/bff/mobile.py::mobile_feed",
                    "elevator_control/adapters/inbound/api/bff/desktop.py::desktop_lift_workbench",
                ],
            ),
            (
                "4.3.1 Domain Events (минимум 2 события)",
                "Реализовано 12 типов доменных событий. Обязательные два — LiftCreated и ServiceRequestCreated, "
                "плюс LiftUpdated/Deleted, ServiceRequestUpdated/Deleted, SensorChanged/Deleted, EventLogged, "
                "TechnicianChanged/Deleted, ReportCreated/Deleted.",
                [
                    "elevator_control/application/events/domain_events.py — все типы и фабрики событий",
                ],
            ),
            (
                "4.3.2 Event-очередь: Command → Event → Queue → Worker",
                "Поток: command публикует событие в outbox (таблица domain_events_log) в той же транзакции, "
                "что и write. Хук Session.after_commit ставит задачу process_domain_event в Celery. Worker "
                "читает событие из outbox и применяет к read-модели через read_sync.sync_*.",
                [
                    "elevator_control/application/events/publisher.py — outbox publish + schedule",
                    "elevator_control/infrastructure/database.py — after_commit hook",
                    "elevator_control/application/events/handlers.py — process_domain_event task",
                    "elevator_control/infrastructure/celery_app.py — регистрация task'а",
                ],
            ),
            (
                "4.4 Eventual Consistency",
                "Команда отдаёт ответ клиенту сразу. Read-модель обновляется через worker с задержкой "
                "CQRS_EVENT_DELAY_SECONDS (по умолчанию 2 сек). Это эмулирует реальный read-after-write lag: "
                "сразу после POST GET по read-модели может ещё не вернуть новую строку.",
                [
                    "elevator_control/application/events/publisher.py — countdown=delay при apply_async",
                    "elevator_control/infrastructure/config.py — cqrs_event_delay_seconds",
                ],
            ),
        ],
    )

    # ---------- Задание 5 ----------
    _add_section(
        doc,
        "Задание 5. Git / Observability / Кэш / Документация",
        [
            (
                "5.1.1 Стратегия веток",
                "Использован feature workflow: каждый блок реализован в отдельной ветке "
                "feature/4.1-cqrs, feature/4.3-events-queue, feature/4.2-bff, feature/6-security, "
                "feature/5.2-observability, docs/readme-and-report. Базовая ветка main защищена.",
                ["git log --oneline --decorate --graph --all"],
            ),
            (
                "5.1.2 Feature workflow",
                "Каждая задача из списка попадала в свою feature-ветку. Изменения, относящиеся к одному заданию, "
                "не попадали в ветку другого. Это видно в git log: каждый коммит имеет conventional-prefix "
                "(feat/docs) и явный номер пункта в заголовке.",
                ["git log --oneline на каждой feature-ветке"],
            ),
            (
                "5.1.3 Pull request — описание изменений",
                "Описания PR оформлены в виде conventional-commit body (что сделано / где смотреть / как проверить) "
                "плюс добавлены отдельные .md файлы в docs/pull_requests/ (см. ниже).",
                ["docs/pull_requests/*.md"],
            ),
            (
                "5.1.4 Стандартизация commit message",
                "Используется Conventional Commits: feat(scope): краткое описание, body с пунктами задания, "
                "трейлер Co-Authored-By. Пример: 'feat(cqrs): 4.1.1, 4.1.2, 4.1.3 — разделение Command/Query, "
                "read-модель, query без ORM'.",
                ["git log --pretty=format:'%h %s' main..HEAD"],
            ),
            (
                "5.2.1 Горячие точки",
                "Три hot-points с метриками: query (GET /api/v1/* и /bff/*), command (POST/PATCH/DELETE), "
                "worker (process_domain_event). Эндпоинт /metrics возвращает агрегаты за последние 60 секунд.",
                [
                    "elevator_control/application/observability.py — record + snapshot",
                    "elevator_control/main.py — RequestLoggingMiddleware вызывает observability.record",
                    "elevator_control/application/events/handlers.py — record для worker",
                ],
            ),
            (
                "5.2.2 Метрики через логи",
                "Каждый hot-point пишет структурированный лог с тегом '5.2.2 metric: ...': время выполнения, "
                "количество строк, ошибки. Кроме того, /metrics отдаёт snapshot (avg_ms / max_ms / calls / errors).",
                [
                    "Логи uvicorn: '5.2.2 metric: HTTP ...' и '5.2.2 metric: query ...'",
                    "Логи celery worker: '5.2.2 metric: worker process_domain_event ...'",
                    "GET /metrics — JSON-снимок",
                ],
            ),
            (
                "5.2.3 Кэш + инвалидация после command",
                "In-process TTL-кэш для query-стороны (TTL 30 сек по умолчанию). Применяется к /api/v1/lifts (list). "
                "Инвалидация: после command сервис вызывает invalidate_for_aggregate; повторно — в воркере после "
                "применения события (на случай распределённых процессов).",
                [
                    "elevator_control/application/cache.py — TTL-кэш с тегами",
                    "elevator_control/application/services.py::_publish_and_invalidate",
                    "elevator_control/application/events/handlers.py — invalidate после apply_event",
                ],
            ),
            (
                "5.3 README",
                "Полный README с разделами: архитектура, инструкции запуска, описание API, явные указания где "
                "смотреть CQRS, где смотреть очередь, где смотреть пункты безопасности.",
                ["README.md"],
            ),
        ],
    )

    # ---------- Задание 6 ----------
    _add_section(
        doc,
        "Задание 6. Безопасность",
        [
            (
                "6.1.1 Валидация входных данных",
                "Все Pydantic-модели имеют extra='forbid' и явные ограничения по длине и типу. "
                "Email проверяется регулярным выражением (без зависимости email-validator), пароль — min_length=8.",
                ["elevator_control/adapters/inbound/api/schemas.py"],
            ),
            (
                "6.1.2 Защита от Mass Assignment",
                "DTO для Create/Update не содержат owner_id, id, created_at. Лишние поля в payload "
                "отвергаются Pydantic с HTTP 422 (extra_forbidden). Проверено: POST с owner_id=99999 → 422.",
                [
                    "elevator_control/adapters/inbound/api/schemas.py — все *Create/*Update классы",
                    "elevator_control/application/services.py — owner_id берётся из current_user/lift",
                ],
            ),
            (
                "6.1.3 Ограничение доступа: owner_id из токена",
                "В Command-сервисах owner_id для новых сущностей всегда выставляется из current_user.id "
                "(пользователь из JWT), а не из тела запроса. Для дочерних сущностей — наследуется от "
                "родительского лифта.",
                [
                    "elevator_control/application/services.py::LiftCommandService.create — owner_id=actor.id",
                    "elevator_control/application/services.py::SensorCommandService.create — owner_id=lift.owner_id",
                ],
            ),
            (
                "6.2.1 Срок жизни access-токена",
                "Короткий TTL access-токена — 15 минут (ACCESS_TOKEN_TTL_SECONDS=900). Каждый JWT содержит "
                "поля iat/exp/jti/type. Просрочка → 401 'Токен истёк'.",
                [
                    "elevator_control/application/auth.py::create_token / decode_token",
                    "elevator_control/infrastructure/config.py::access_token_ttl_seconds",
                ],
            ),
            (
                "6.2.2 Refresh-токен и его обновление",
                "При login отдаётся пара access+refresh. Refresh имеет TTL 7 суток. Эндпоинт POST /auth/refresh "
                "принимает refresh_token и выдаёт НОВУЮ пару, при этом старый refresh-jti помещается в blacklist "
                "(rotation). Это предотвращает повторное использование refresh-токена.",
                [
                    "elevator_control/adapters/inbound/api/v1/auth.py::refresh_token",
                    "elevator_control/application/auth.py::AuthApplicationService.refresh",
                ],
            ),
            (
                "6.2.3 Logout / blacklist",
                "Эндпоинт POST /auth/logout помещает jti текущего access-токена в таблицу revoked_tokens. "
                "При следующем запросе с этим же токеном — 401 'Токен отозван'. Refresh-rotation использует "
                "тот же механизм. Запись хранится до естественного истечения TTL токена (expires_at).",
                [
                    "alembic/versions/003_cqrs_read_model.py — таблица revoked_tokens",
                    "elevator_control/application/auth.py::_revoke / _is_revoked",
                    "elevator_control/adapters/inbound/api/v1/auth.py::logout",
                ],
            ),
            (
                "6.3.1 Rate limiting",
                "Middleware с алгоритмом скользящего окна per-IP. Два предела: rate_limit_per_minute и "
                "rate_limit_burst_per_10s. Превышение → 429 Too Many Requests с Retry-After. Не лимитирует "
                "preflight CORS и /health.",
                ["elevator_control/main.py::RateLimitMiddleware"],
            ),
            (
                "6.3.2 SQL-инъекции",
                "Все запросы используют параметризацию: SQLAlchemy text(...:name) и asyncpg ($1, $2). "
                "В Query-сервисах WHERE собирается из доверенных имён колонок, значения никогда не "
                "интерполируются в SQL-текст. Демонстрация: GET /api/v1/lifts/1%20OR%201=1 → 422 (ошибка типа).",
                [
                    "elevator_control/application/queries/*.py",
                    "elevator_control/application/read_sync.py",
                    "elevator_control/application/auth.py — _is_revoked / _revoke",
                ],
            ),
            (
                "6.3.3 XSS / security headers",
                "Глобальный SecurityHeadersMiddleware проставляет X-Content-Type-Options=nosniff, X-Frame-Options=DENY, "
                "Referrer-Policy=no-referrer и Content-Security-Policy. Для /api и /bff: default-src 'none' "
                "(JSON-ответ браузер не исполняет как HTML); для статических /clients разрешены self+inline.",
                ["elevator_control/main.py::SecurityHeadersMiddleware"],
            ),
            (
                "6.3.4 CORS только разрешённые домены",
                "CORSMiddleware ограничен whitelist'ом из CORS_ALLOWED_ORIGINS (никаких '*'). Демонстрация: "
                "preflight с http://localhost:3000 → 200; preflight с http://evil.example.com → 400 Bad Request.",
                [
                    "elevator_control/infrastructure/config.py::cors_allowed_origins",
                    "elevator_control/main.py — app.add_middleware(CORSMiddleware, allow_origins=...)",
                ],
            ),
        ],
    )

    # ---------- Проверка ----------
    _add_heading(doc, "Как проверить вживую", level=1)
    doc.add_paragraph(
        "1) Применить миграции: alembic upgrade head\n"
        "2) Запустить API: python -m uvicorn elevator_control.main:app --reload --port 8000\n"
        "3) Запустить воркер: celery -A elevator_control.infrastructure.celery_app worker --loglevel=info --pool=solo\n"
        "4) Зарегистрировать первого администратора (он получит роль administrator автоматически):\n"
        "   curl -X POST http://localhost:8000/api/v1/auth/register -H 'Content-Type: application/json' \\\n"
        "        -d '{\"email\":\"admin@local\",\"password\":\"password123\",\"role\":\"administrator\"}'\n"
        "5) Войти: curl -X POST http://localhost:8000/api/v1/auth/login-json -H 'Content-Type: application/json' \\\n"
        "        -d '{\"email\":\"admin@local\",\"password\":\"password123\"}'\n"
        "6) Дальше — токен передавать в заголовке Authorization: Bearer <access_token>.\n"
        "7) Создать лифт: POST /api/v1/lifts; через ~2 секунды read-модель догонит — GET /api/v1/lifts.\n"
        "8) Открыть /docs (Swagger) для интерактивной проверки. /metrics — снимок hot-points.\n"
    )

    out_path = Path(__file__).resolve().parent / "report.docx"
    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
